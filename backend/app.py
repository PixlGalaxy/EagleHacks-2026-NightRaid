from flask import Flask, jsonify, request, Response, stream_with_context
from flask_cors import CORS
from datetime import datetime
import psutil
import os
import requests
import json
import re
import logging
from logging.handlers import RotatingFileHandler
from werkzeug.utils import secure_filename
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False

app = Flask(__name__)

################################################################################################################
## LOGGING CONFIGURATION

# Determine log level based on environment
FLASK_ENV = os.environ.get('FLASK_ENV', 'production')
LOG_LEVEL = logging.DEBUG if FLASK_ENV == 'development' else logging.INFO

# Create logs directory if it doesn't exist
LOGS_DIR = os.path.join(os.path.dirname(__file__), 'logs')
os.makedirs(LOGS_DIR, exist_ok=True)

# Configure logging formatter
log_formatter = logging.Formatter(
    '[%(asctime)s] %(levelname)-8s | %(name)s | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# Console handler (always enabled)
console_handler = logging.StreamHandler()
console_handler.setLevel(LOG_LEVEL)
console_handler.setFormatter(log_formatter)

# File handler (rotating files)
file_handler = RotatingFileHandler(
    os.path.join(LOGS_DIR, 'nightraid.log'),
    maxBytes=10*1024*1024,  # 10MB per file
    backupCount=5  # Keep 5 backup files
)
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(log_formatter)

# Configure root logger
root_logger = logging.getLogger()
root_logger.setLevel(LOG_LEVEL)
root_logger.addHandler(console_handler)
root_logger.addHandler(file_handler)

# Configure Flask logger
logging.getLogger('werkzeug').setLevel(LOG_LEVEL)

# Create app logger
logger = logging.getLogger(__name__)
logger.info("="*80)
logger.info("NightRaid Banking Analysis Backend - Logger Initialized")
logger.info("="*80)

################################################################################################################
## CORS CONFIGURATION
CORS_ORIGINS = os.environ.get('CORS_ORIGINS', '*').split(',')
cors_config = {
    "origins": CORS_ORIGINS,
    "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization"],
    "supports_credentials": True,
    "max_age": 3600
}
CORS(app, resources={r"/api/*": cors_config, r"/uptime": cors_config, r"/health": cors_config, r"/main/*": cors_config})

################################################################################################################
## GLOBAL VARIABLES AND CONFIGURATION AND DEFINITIONS
server_start_time = datetime.now()

# Flask Configuration
FLASK_ENV = os.environ.get('FLASK_ENV', 'production')
FLASK_DEBUG = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
SERVER_PORT = int(os.environ.get('PORT', 5000))

# Ollama Configuration (loaded from .env)
OLLAMA_BASE_URL = os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')
OLLAMA_MODEL = os.environ.get('OLLAMA_MODEL', 'llama2')
OLLAMA_API_KEY = os.environ.get('OLLAMA_API_KEY', '')
OLLAMA_TIMEOUT = int(os.environ.get('OLLAMA_TIMEOUT', 120))  # seconds

# File Upload Configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
ALLOWED_EXTENSIONS = {'txt', 'csv', 'pdf', 'json', 'xlsx', 'xls', 'docx', 'doc'}
MAX_FILE_SIZE = int(os.environ.get('MAX_FILE_SIZE', 10 * 1024 * 1024))  # 10MB maximum file size
MAX_PROMPT_CHARS = int(os.environ.get('MAX_PROMPT_CHARS', 50000))  # max chars of file content sent to model

# Create uploads folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

################################################################################################################
## UTILITY FUNCTIONS FOR FILE PROCESSING AND ANALYSIS

def allowed_file(filename):
    """
    Validate if uploaded file has an allowed extension.
    Allowed formats: txt, csv, pdf, json, xlsx, xls, docx, doc
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_uploaded_file(file):
    """
    Save uploaded file to the uploads folder.
    Returns filepath if successful, None otherwise.
    """
    if not file or file.filename == '':
        return None
    
    if not allowed_file(file.filename):
        return None
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    return filepath

def read_file_content(filepath):
    """
    Read and return file content based on file extension.
    Supports: TXT, CSV, JSON, PDF, XLSX, XLS, DOCX, DOC
    """
    try:
        if filepath.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        elif filepath.endswith('.csv'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        elif filepath.endswith('.json'):
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        elif filepath.endswith('.pdf'):
            try:
                import PyPDF2
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                    return text
            except ImportError:
                return "[PDF file detected - PDF parsing library not available. Please install PyPDF2]"
        elif filepath.endswith('.xlsx'):
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(filepath)
                text = ""
                for sheet in workbook.sheetnames:
                    text += f"\n=== Sheet: {sheet} ===\n"
                    ws = workbook[sheet]
                    for row in ws.iter_rows(values_only=True):
                        text += " | ".join([str(cell) if cell is not None else "" for cell in row]) + "\n"
                return text
            except ImportError:
                return "[XLSX file detected - Excel library not available. Please install openpyxl]"
        elif filepath.endswith('.xls'):
            try:
                import xlrd
                workbook = xlrd.open_workbook(filepath)
                text = ""
                for sheet_idx in range(workbook.nsheets):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    text += f"\n=== Sheet: {sheet.name} ===\n"
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                return text
            except ImportError:
                return "[XLS file detected - Excel library not available. Please install xlrd]"
        elif filepath.endswith('.docx'):
            try:
                from docx import Document
                doc = Document(filepath)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    text += "\n=== Table ===\n"
                    for row in table.rows:
                        text += " | ".join([cell.text for cell in row.cells]) + "\n"
                return text
            except ImportError:
                return "[DOCX file detected - Word library not available. Please install python-docx]"
        elif filepath.endswith('.doc'):
            try:
                import docx2txt
                text = docx2txt.process(filepath)
                return text
            except ImportError:
                return "[DOC file detected - Word library not available. Please install python-docx2txt]"
        return None
    except Exception as e:
        return f"Error reading file: {str(e)}"

def extract_bank_statement_info(content):
    """
    Extract key information from bank statement content using regex patterns.
    Returns: account holder, account number, statement period, transaction count
    """
    info = {
        'account_holder': None,
        'account_number': None,
        'statement_period': None,
        'total_transactions': 0,
        'preview': content[:500] if content else ""
    }
    
    # Extract account number (support multiple formats)
    account_match = re.search(r'(?:Account|Cuenta)\s*(?:Number|Números?|#)?[\s:]*([A-Za-z0-9\-*]{10,})', content, re.IGNORECASE)
    if account_match:
        info['account_number'] = account_match.group(1)
    
    # Extract account holder name
    holder_match = re.search(r'(?:Account\s*(?:Holder|Owner)|Nombre|Name)[\s:]*([A-Za-z\s]+)', content, re.IGNORECASE)
    if holder_match:
        info['account_holder'] = holder_match.group(1).strip()
    
    # Extract statement period
    period_match = re.search(r'(?:Statement|Periodo|Period)[\s:]*([A-Za-z0-9\s,/-]+)', content, re.IGNORECASE)
    if period_match:
        info['statement_period'] = period_match.group(1).strip()
    
    # Count estimated transactions by looking for date patterns
    transaction_lines = len(re.findall(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', content))
    info['total_transactions'] = transaction_lines
    
    return info

def generate_bank_statement_prompt(analysis_type, file_content, file_metadata):
    """
    Generate context-aware prompts for Ollama based on analysis type.
    Includes a system message task definition and bank statement content.
    """
    
    bank_info = extract_bank_statement_info(file_content)

    # System-level task instructions per analysis type | for AI model to understand the specific analysis focus
    system_messages = {
        'spending': (
            "You are an expert financial analyst. Your task is to analyze the SPENDING BEHAVIOR "
            "in this bank statement. Break down spending by category, identify trends, and highlight "
            "the top expense areas. Use **bold** for key figures, bullet points for categories, "
            "and Markdown tables where helpful."
        ),
        'anomalies': (
            "You are a fraud detection specialist. Your task is to DETECT ANOMALIES AND SUSPICIOUS "
            "TRANSACTIONS in this bank statement. Flag unusual amounts, irregular patterns, and "
            "assign a risk level (Low / Medium / High) to each finding. Use **bold** for flagged "
            "items and a Markdown table to summarize anomalies."
        ),
        'risk': (
            "You are a financial risk advisor. Your task is to perform a FINANCIAL RISK ASSESSMENT "
            "of this bank statement. Provide an overall Risk Score (1–10), evaluate Liquidity Risk, "
            "Fraud Risk, and Income Volatility Risk. Use a Markdown table for the risk breakdown and "
            "**bold** for scores and critical warnings."
        ),
        'recommendations': (
            "You are a personal finance coach. Your task is to generate ACTIONABLE FINANCIAL "
            "RECOMMENDATIONS based on this bank statement. Prioritize by impact and feasibility. "
            "Use numbered lists for action items, **bold** for key advice, and Markdown tables "
            "to compare options where relevant."
        ),
        'summary': (
            "You are a certified financial analyst. Your task is to produce a COMPREHENSIVE "
            "FINANCIAL SUMMARY of this bank statement. Cover financial health, income vs expenses, "
            "key metrics, major trends, and top 3 priority actions. Use Markdown headers (##), "
            "tables, bullet points, and **bold** for important figures."
        ),
    }

    system_msg = system_messages.get(analysis_type, system_messages['summary'])

    base_prompt = f"""{system_msg}

---

**BANK STATEMENT DETAILS**
- Account Holder: {bank_info['account_holder'] or 'Unknown'}
- Account Number: {bank_info['account_number'] or 'Unknown'}
- Statement Period: {bank_info['statement_period'] or 'Unknown'}
- Estimated Transactions: {bank_info['total_transactions']}

**STATEMENT CONTENT**
```
{file_content[:MAX_PROMPT_CHARS]}
```
"""

    analysis_instructions = {
        'spending': f"""{base_prompt}

## Spending Behavior Analysis

Provide:
1. **Total spending** and breakdown by category (use a table with Category | Amount | % of Total)
2. **Top 5 expense categories** with specific amounts
3. **Spending trends** and patterns (weekly/monthly if detectable)
4. **Notable observations** about financial behavior""",

        'anomalies': f"""{base_prompt}

## Anomaly & Fraud Detection

Provide:
1. A **summary table** of flagged transactions (Date | Description | Amount | Risk Level)
2. **Explanation** of why each transaction is flagged
3. **Overall fraud risk** assessment (Low / Medium / High)
4. **Recommended actions** for each anomaly""",

        'risk': f"""{base_prompt}

## Financial Risk Assessment

Provide:
1. **Overall Risk Score** (1–10) with justification
2. A **risk breakdown table** (Risk Category | Score | Details)
   - Liquidity Risk
   - Fraud Risk
   - Income Volatility Risk
3. **Red flags** detected
4. **Mitigation recommendations**""",

        'recommendations': f"""{base_prompt}

## Financial Recommendations

Provide:
1. **Top savings opportunities** with estimated savings amounts
2. **Budget optimization** suggestions
3. **Investment or wealth-building** opportunities
4. **30-60-90 day action plan** as a prioritized list""",

        'summary': f"""{base_prompt}

## Comprehensive Financial Summary

Provide:
1. **Financial Health Overview** (1 paragraph executive summary)
2. **Income vs Expenses table** (Category | Amount | Notes)
3. **Key Financial Metrics** (table format)
4. **Major Trends** identified
5. **Top 3 Priority Actions** (numbered, with expected impact)""",
    }

    return analysis_instructions.get(analysis_type, analysis_instructions['summary'])

def call_ollama_with_prompt(prompt, model=None):
    """
    Call Ollama API with the given prompt and return the generated response.
    Handles connection errors and timeouts gracefully.
    """
    if not model:
        model = OLLAMA_MODEL
    
    try:
        response = requests.post(
            f'{OLLAMA_BASE_URL}/api/generate',
            json={
                'model': model,
                'prompt': prompt,
                'stream': False
            },
            timeout=OLLAMA_TIMEOUT
        )
        
        if response.status_code == 200:
            return response.json().get('response', '')
        else:
            raise Exception(f"Ollama API error: {response.text}")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Could not connect to Ollama at {OLLAMA_BASE_URL}")
    except requests.exceptions.Timeout:
        raise Exception("Ollama request timeout - model processing took too long")

################################################################################################################
## API ROUTES

@app.route('/uptime', methods=['GET'])
def uptime():
    """
    Endpoint: GET /uptime
    Returns server uptime status, system metrics, and service information.
    Response: JSON with uptime formatted and system resource usage.
    """
    logger.debug("GET /uptime - Request received")
    try:
        current_time = datetime.now()
        uptime_delta = current_time - server_start_time
        
        total_seconds = int(uptime_delta.total_seconds())
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        seconds = total_seconds % 60
        
        cpu_percent = psutil.cpu_percent(interval=1)
        memory_info = psutil.virtual_memory()
        
        response = {
            'status': 'online',
            'timestamp': current_time.isoformat(),
            'server_start_time': server_start_time.isoformat(),
            'uptime': {
                'total_seconds': total_seconds,
                'formatted': f'{hours}h {minutes}m {seconds}s'
            },
            'system': {
                'cpu_usage_percent': round(cpu_percent, 2),
                'memory_usage_percent': round(memory_info.percent, 2),
                'memory_available_mb': round(memory_info.available / (1024 * 1024), 2)
            },
            'service': 'NightRaid Banking Analysis Backend API'
        }
        
        return jsonify(response), 200
    
    except Exception as e:
        logger.error(f"GET /uptime - Error: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/health', methods=['GET'])
def health():
    """
    Endpoint: GET /health
    Simple health check endpoint for monitoring and load balancers.
    Response: JSON status indicating service availability.
    """
    return jsonify({
        'status': 'healthy',
        'service': 'Running',
        'timestamp': datetime.now().isoformat()
    }), 200


@app.route('/', methods=['GET'])
def index():
    """
    Endpoint: GET /
    Root endpoint providing API documentation and available endpoints.
    Response: JSON object documenting all available API routes.
    """
    return jsonify({
        'message': 'Welcome to NightRaid Banking Analysis Backend API',
        'version': '0.3.0',
        'description': 'AI-powered banking analytics and financial insights using Ollama',
        'endpoints': {
            'system': {
                'description': 'System health and monitoring',
                'uptime': 'GET /uptime',
                'health': 'GET /health'
            },
            'ollama': {
                'description': 'Direct Ollama model access',
                'generate': 'POST /ollama/generate',
                'models': 'GET /ollama/models'
            },
            'banking_analysis': {
                'description': 'Bank statement analysis and processing',
                'upload': 'POST /main/upload',
                'analyze': 'POST /main/analyze',
                'analyze_spending': 'POST /main/analyze/spending',
                'analyze_anomalies': 'POST /main/analyze/anomalies',
                'analyze_risk': 'POST /main/analyze/risk',
                'analyze_recommendations': 'POST /main/analyze/recommendations',
                'analyze_summary': 'POST /main/analyze/summary'
            }
        },
        'documentation': 'See endpoint descriptions in docstrings'
    }), 200


@app.route('/ollama/generate', methods=['POST'])
def ollama_generate():
    """
    Endpoint: POST /ollama/generate
    Direct access to Ollama text generation.
    Request body: {'prompt': 'text to process', 'model': 'optional_model_name'}
    Response: Generated text from Ollama model.
    """
    if not OLLAMA_AVAILABLE:
        return jsonify({
            'status': 'error',
            'message': 'Ollama library not installed'
        }), 503
    
    try:
        data = request.get_json()
        if not data or 'prompt' not in data:
            return jsonify({
                'status': 'error',
                'message': 'Missing required field: prompt'
            }), 400
        
        prompt = data.get('prompt')
        model = data.get('model', OLLAMA_MODEL)
        
        response = requests.post(
            f'{OLLAMA_BASE_URL}/api/generate',
            json={
                'model': model,
                'prompt': prompt,
                'stream': False
            },
            timeout=OLLAMA_TIMEOUT
        )
        
        if response.status_code != 200:
            return jsonify({
                'status': 'error',
                'message': f'Ollama API error: {response.text}'
            }), response.status_code
        
        result = response.json()
        
        return jsonify({
            'status': 'success',
            'model': model,
            'prompt': prompt,
            'response': result.get('response', ''),
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except requests.exceptions.ConnectionError:
        return jsonify({
            'status': 'error',
            'message': f'Could not connect to Ollama at {OLLAMA_BASE_URL}. Make sure Ollama is running.'
        }), 503
    except requests.exceptions.Timeout:
        return jsonify({
            'status': 'error',
            'message': 'Ollama request timeout. The model may be processing too long.'
        }), 504
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@app.route('/main/upload', methods=['POST'])
def upload_files():
    """
    Endpoint: POST /main/upload
    Upload bank statement files for analysis.
    Accepted formats: CSV, TXT, JSON, PDF, XLSX, XLS, DOCX, DOC
    Max file size: 10MB per file
    
    Request: Multipart form data with files (file0, file1, etc. or files field)
    Response: JSON with uploaded file metadata (filename, path, size, type)
    """
    logger.info("POST /main/upload - File upload request received")
    try:
        # Validate request contains files
        if 'files' not in request.files and len(request.files) == 0:
            logger.warning("Upload failed: No files field in request")
            return jsonify({
                'status': 'error',
                'message': 'No files field found in request'
            }), 400
        
        # Extract files from request (handle multiple naming conventions)
        files = request.files.getlist('file0') if 'file0' in request.files else request.files.getlist('files')
        
        # Fallback: handle file0, file1, file2, etc. format
        if not files or (len(files) == 1 and files[0].filename == ''):
            files = []
            index = 0
            while f'file{index}' in request.files:
                file_obj = request.files.get(f'file{index}')
                if file_obj:
                    files.append(file_obj)
                index += 1
        
        if not files:
            logger.warning("Upload failed: No files extracted from request")
            return jsonify({
                'status': 'error',
                'message': 'No files provided or request format incorrect'
            }), 400
        
        logger.info(f"Processing {len(files)} file(s)...")
        uploaded_info = []
        
        # Process each uploaded file
        for idx, file in enumerate(files, 1):
            if not file or file.filename == '':
                logger.debug(f"File {idx}: Skipped (empty filename)")
                continue
            
            # Validate file extension
            if not allowed_file(file.filename):
                logger.debug(f"File {idx}: '{file.filename}' - Invalid extension")
                continue
            
            logger.debug(f"  File {idx}: '{file.filename}' ({file.content_type})")
            # Save file to uploads folder
            filepath = save_uploaded_file(file)
            
            if filepath and os.path.exists(filepath):
                file_size = os.path.getsize(filepath)
                logger.info(f"File {idx}: '{file.filename}' saved ({file_size} bytes)")
                uploaded_info.append({
                    'filename': file.filename,
                    'filepath': filepath,
                    'size': file_size,
                    'type': file.content_type,
                    'upload_time': datetime.now().isoformat()
                })
            else:
                logger.error(f"File {idx}: Failed to save '{file.filename}'")
        
        if not uploaded_info:
            logger.error("Upload failed: No valid files processed")
            return jsonify({
                'status': 'error',
                'message': 'No valid files uploaded. Allowed formats: .txt, .csv, .json, .pdf'
            }), 400
        
        logger.info(f"Upload successful: {len(uploaded_info)} file(s) processed")
        return jsonify({
            'status': 'success',
            'message': f'{len(uploaded_info)} file(s) uploaded successfully',
            'files': uploaded_info,
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        logger.error(f"POST /main/upload - Exception: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/main/analyze', methods=['POST'])
def analyze_bank_statement():
    """
    Endpoint: POST /main/analyze
    Process bank statements through Ollama AI for financial analysis.
    
    Request body:
    {
        'analysisType': 'spending|anomalies|risk|recommendations|summary',
        'files': [{'name': 'filename.csv', 'type': 'text/csv', 'size': 1024}, ...]
    }
    
    Response: JSON array with analysis results per file
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'Request body must contain valid JSON'
            }), 400
        
        # Extract analysis type and files metadata
        analysis_type = data.get('analysisType', 'summary')
        files_metadata = data.get('files', [])
        
        # Validate input
        if not files_metadata:
            return jsonify({
                'status': 'error',
                'message': 'No files metadata provided in request body'
            }), 400
        
        # Validate analysis type
        valid_types = ['spending', 'anomalies', 'risk', 'recommendations', 'summary']
        if analysis_type not in valid_types:
            return jsonify({
                'status': 'error',
                'message': f'Invalid analysis type. Allowed: {valid_types}'
            }), 400
        
        results = []
        
        # Process each file
        for file_meta in files_metadata:
            filename = file_meta.get('name')
            if not filename:
                results.append({
                    'status': 'error',
                    'message': 'File metadata missing name field'
                })
                continue
            
            # Build file path
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
            
            # Check if file exists
            if not os.path.exists(filepath):
                results.append({
                    'filename': filename,
                    'status': 'error',
                    'message': f'File not found in uploads folder: {filename}'
                })
                continue
            
            try:
                # Read and validate file content
                file_content = read_file_content(filepath)
                
                if not file_content or file_content.startswith('Error'):
                    results.append({
                        'filename': filename,
                        'status': 'error',
                        'message': file_content or 'Unable to read file content'
                    })
                    continue
                
                # Generate analysis prompt based on type
                prompt = generate_bank_statement_prompt(analysis_type, file_content, file_meta)
                
                # Send to Ollama for analysis
                analysis_result = call_ollama_with_prompt(prompt)
                
                if not analysis_result:
                    results.append({
                        'filename': filename,
                        'status': 'error',
                        'message': 'Ollama returned empty response'
                    })
                    continue
                
                # Add successful result
                results.append({
                    'filename': filename,
                    'status': 'success',
                    'analysis_type': analysis_type,
                    'analysis': analysis_result,
                    'timestamp': datetime.now().isoformat()
                })
            
            except Exception as e:
                results.append({
                    'filename': filename,
                    'status': 'error',
                    'error_type': type(e).__name__,
                    'message': str(e)
                })
        
        return jsonify({
            'status': 'success',
            'analysis_type': analysis_type,
            'results': results,
            'timestamp': datetime.now().isoformat()
        }), 200
    
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/main/analyze/stream', methods=['POST'])
def analyze_stream():
    """
    Endpoint: POST /main/analyze/stream
    Streaming version of the analysis endpoint using Server-Sent Events (SSE).
    Tokens are streamed back incrementally as Ollama generates them.

    Request body: {'analysisType': '...', 'files': [...]}
    Response: text/event-stream with JSON data payloads per chunk
    """
    data = request.get_json()

    if not data:
        return jsonify({'status': 'error', 'message': 'Request body must contain valid JSON'}), 400

    analysis_type = data.get('analysisType', 'summary')
    files_metadata = data.get('files', [])

    valid_types = ['spending', 'anomalies', 'risk', 'recommendations', 'summary']
    if analysis_type not in valid_types:
        return jsonify({'status': 'error', 'message': f'Invalid analysis type. Allowed: {valid_types}'}), 400

    if not files_metadata:
        return jsonify({'status': 'error', 'message': 'No files metadata provided'}), 400

    def generate():
        for file_meta in files_metadata:
            filename = file_meta.get('name')
            if not filename:
                yield f"data: {json.dumps({'type': 'error', 'message': 'File metadata missing name field'})}\n\n"
                continue

            filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))

            if not os.path.exists(filepath):
                yield f"data: {json.dumps({'type': 'error', 'filename': filename, 'message': f'File not found: {filename}'})}\n\n"
                continue

            file_content = read_file_content(filepath)
            if not file_content or file_content.startswith('Error'):
                yield f"data: {json.dumps({'type': 'error', 'filename': filename, 'message': file_content or 'Unable to read file'})}\n\n"
                continue

            prompt = generate_bank_statement_prompt(analysis_type, file_content, file_meta)

            # Notify frontend a new file analysis is starting
            yield f"data: {json.dumps({'type': 'file_start', 'filename': filename, 'analysis_type': analysis_type})}\n\n"

            try:
                ollama_res = requests.post(
                    f'{OLLAMA_BASE_URL}/api/generate',
                    json={'model': OLLAMA_MODEL, 'prompt': prompt, 'stream': True},
                    stream=True,
                    timeout=OLLAMA_TIMEOUT
                )

                for line in ollama_res.iter_lines():
                    if not line:
                        continue
                    try:
                        chunk = json.loads(line)
                        token = chunk.get('response', '')
                        if token:
                            yield f"data: {json.dumps({'type': 'token', 'text': token})}\n\n"
                        if chunk.get('done'):
                            break
                    except json.JSONDecodeError:
                        continue

                yield f"data: {json.dumps({'type': 'file_end', 'filename': filename})}\n\n"

            except requests.exceptions.ConnectionError:
                yield f"data: {json.dumps({'type': 'error', 'filename': filename, 'message': f'Could not connect to Ollama at {OLLAMA_BASE_URL}'})}\n\n"
            except requests.exceptions.Timeout:
                yield f"data: {json.dumps({'type': 'error', 'filename': filename, 'message': 'Ollama request timed out'})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'filename': filename, 'message': str(e)})}\n\n"

        yield "data: [DONE]\n\n"

    return Response(
        stream_with_context(generate()),
        content_type='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'close',
        }
    )


@app.route('/main/analyze/spending', methods=['POST'])
def analyze_spending():
    """
    Endpoint: POST /main/analyze/spending
    Specialized endpoint for spending behavior analysis.
    Request body: {'files': [file_metadata]}
    Automatically sets analysisType to 'spending'.
    """
    data = request.get_json() or {}
    data['analysisType'] = 'spending'
    request.json = data
    return analyze_bank_statement()


@app.route('/main/analyze/anomalies', methods=['POST'])
def analyze_anomalies():
    """
    Endpoint: POST /main/analyze/anomalies
    Specialized endpoint for anomaly and fraud detection.
    Request body: {'files': [file_metadata]}
    Automatically sets analysisType to 'anomalies'.
    """
    data = request.get_json() or {}
    data['analysisType'] = 'anomalies'
    request.json = data
    return analyze_bank_statement()


@app.route('/main/analyze/risk', methods=['POST'])
def analyze_risk():
    """
    Endpoint: POST /main/analyze/risk
    Specialized endpoint for financial risk assessment.
    Request body: {'files': [file_metadata]}
    Automatically sets analysisType to 'risk' and returns risk scores.
    """
    data = request.get_json() or {}
    data['analysisType'] = 'risk'
    request.json = data
    return analyze_bank_statement()


@app.route('/main/analyze/recommendations', methods=['POST'])
def analyze_recommendations():
    """
    Endpoint: POST /main/analyze/recommendations
    Specialized endpoint for financial recommendations.
    Request body: {'files': [file_metadata]}
    Automatically sets analysisType to 'recommendations'.
    """
    data = request.get_json() or {}
    data['analysisType'] = 'recommendations'
    request.json = data
    return analyze_bank_statement()


@app.route('/main/analyze/summary', methods=['POST'])
def analyze_summary():
    """
    Endpoint: POST /main/analyze/summary
    Specialized endpoint for comprehensive financial summary.
    Request body: {'files': [file_metadata]}
    Automatically sets analysisType to 'summary'.
    """
    data = request.get_json() or {}
    data['analysisType'] = 'summary'
    request.json = data
    return analyze_bank_statement()

################################################################################################################
## APPLICATION STARTUP

if __name__ == '__main__':
    print(f"""\n
╔═══════════════════════════════════════════════════════════════╗
║        NightRaid Banking Analysis Backend API                 ║
║           AI-Powered Financial Intelligence                   ║
╚═══════════════════════════════════════════════════════════════╝

Server Configuration (from .env):
Environment: {FLASK_ENV}
Host: 0.0.0.0
Port: {SERVER_PORT}
Debug Mode: {FLASK_DEBUG}
  
Ollama Configuration:
Model: {OLLAMA_MODEL}
Base URL: {OLLAMA_BASE_URL}
Timeout: {OLLAMA_TIMEOUT}s
API Key: {'***' if OLLAMA_API_KEY else 'Not set'}
  
File Management: Upload Folder: {UPLOAD_FOLDER}
Max File Size: {MAX_FILE_SIZE / (1024*1024):.1f}MB
Allowed Extensions: {', '.join(ALLOWED_EXTENSIONS)}
  
CORS Configuration:
Allowed Origins: {', '.join(CORS_ORIGINS)}

Starting server...\n\n""")
app.run(host='0.0.0.0', port=SERVER_PORT, debug=FLASK_DEBUG)
